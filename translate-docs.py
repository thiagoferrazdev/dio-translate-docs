# -*- coding: utf-8 -*-

import os
from os import path
import requests
from docx import Document

# Configurações da API
API_KEY = "API_KEY"  # Substitua pela sua chave de API
ENDPOINT_URL = "ENDPOINT_URL"  # Substitua pela URL do endpoint
LOCATION = "eastus2"  # Região do serviço
TARGET_LANGUAGE = "pt-br"  # Idioma de destino

def translate_text(text, target_language=TARGET_LANGUAGE):
    """
    Traduz um texto para o idioma de destino usando a API de tradução.

    Args:
        text (str): Texto a ser traduzido.
        target_language (str): Idioma de destino (default: pt-br).

    Returns:
        str: Texto traduzido.
    """
    constructed_url = f"{ENDPOINT_URL}/translate"
    headers = {
        "Ocp-Apim-Subscription-Key": API_KEY,
        "Ocp-Apim-Subscription-Region": LOCATION,
        "Content-type": "application/json",
        "X-ClientTraceId": str(os.urandom(16))
    }
    body = [{"text": text}]
    params = {
        "api-version": "3.0",
        "from": "en",
        "to": target_language
    }

    try:
        response = requests.post(constructed_url, params=params, headers=headers, json=body)
        response.raise_for_status()
        response_json = response.json()
        return response_json[0]["translations"][0]["text"]
    except requests.exceptions.RequestException as e:
        raise Exception(f"Erro na API de tradução: {e}")

def translate_document(input_path, output_language=TARGET_LANGUAGE):
    """
    Traduz um documento Word (.docx) para o idioma especificado.

    Args:
        input_path (str): Caminho do arquivo de entrada.
        output_language (str): Idioma de destino (default: pt-br).

    Returns:
        str: Caminho do arquivo traduzido.
    """
    if not path.isfile(input_path) or not input_path.endswith(".docx"):
        raise FileNotFoundError(f"O arquivo {input_path} não é válido ou não foi encontrado.")
    
    try:
        document = Document(input_path)
        translated_doc = Document()

        for paragraph in document.paragraphs:
            if paragraph.text.strip():  # Ignorar parágrafos vazios
                translated_text = translate_text(paragraph.text, output_language)
                translated_doc.add_paragraph(translated_text)

        output_path = input_path.replace(".docx", f"_{output_language}.docx")
        translated_doc.save(output_path)
        print(f"Documento traduzido salvo em: {output_path}")
        return output_path

    except Exception as e:
        raise Exception(f"Erro ao traduzir o documento: {e}")

if __name__ == "__main__":
    # Caminho do arquivo a ser traduzido
    input_file = "/content/Bones.docx"

    try:
        output_file = translate_document(input_file)
        print(f"Tradução concluída com sucesso: {output_file}")
    except Exception as error:
        print(f"Erro durante a execução: {error}")
